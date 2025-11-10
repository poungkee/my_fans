#!/usr/bin/env python3
"""
Markdown to Word Converter
FANS 프로젝트 평가 보고서를 Word 문서로 변환
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def add_heading(doc, text, level=1):
    """제목 추가"""
    heading = doc.add_heading(text, level=level)
    # 한글 폰트 설정
    for run in heading.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(20 - level * 2)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        # 영문 폰트
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return heading

def add_paragraph(doc, text, bold=False, size=11):
    """문단 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    # 한글 폰트
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def add_table_from_markdown(doc, lines):
    """마크다운 테이블을 Word 테이블로 변환"""
    # 테이블 파싱
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|---'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)

    if not rows:
        return None

    # 테이블 생성
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    # 데이터 채우기
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            # 헤더 행 (첫 번째 행) 볼드 처리
            if i == 0:
                run = cell.paragraphs[0].add_run(cell_data)
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run = cell.paragraphs[0].add_run(cell_data)
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    return table

def add_code_block(doc, text):
    """코드 블록 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)
    # 배경색 (회색)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def convert_markdown_to_word(md_file, output_file):
    """마크다운 파일을 Word로 변환"""
    doc = Document()

    # 문서 속성 설정
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # 코드 블록 시작/종료
        if line.startswith('```'):
            if in_code_block:
                # 코드 블록 종료
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # 코드 블록 시작
                in_code_block = True
            i += 1
            continue

        # 코드 블록 내부
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 테이블 감지
        if '|' in line and not line.strip().startswith('#'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            # 다음 줄 확인
            if i < len(lines) and '|' not in lines[i]:
                # 테이블 종료
                add_table_from_markdown(doc, table_lines)
                doc.add_paragraph()  # 여백
                in_table = False
                table_lines = []
            continue

        # 제목
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            # 이모지 및 특수 문자 제거
            text = re.sub(r'[📋📑🔟1️⃣2️⃣3️⃣4️⃣5️⃣6️⃣7️⃣8️⃣9️⃣✅❌]', '', text).strip()
            if text and text != '---':
                add_heading(doc, text, level=min(level, 3))

        # 굵은 글씨 (볼드)
        elif line.startswith('**') and line.endswith('**'):
            text = line.strip('*').strip()
            add_paragraph(doc, text, bold=True, size=12)

        # 리스트
        elif line.startswith('- ') or line.startswith('* '):
            text = line.lstrip('-* ').strip()
            # 볼드 처리
            if text.startswith('**') and '**:' in text:
                parts = text.split('**:')
                bold_part = parts[0].strip('*')
                rest = parts[1] if len(parts) > 1 else ''
                p = doc.add_paragraph(style='List Bullet')
                run1 = p.add_run(bold_part + ':')
                run1.font.bold = True
                run1.font.name = '맑은 고딕'
                run1.font.size = Pt(11)
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                if rest:
                    run2 = p.add_run(rest)
                    run2.font.name = '맑은 고딕'
                    run2.font.size = Pt(11)
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            else:
                p = doc.add_paragraph(text, style='List Bullet')
                for run in p.runs:
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # 번호 리스트
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                run.font.name = '맑은 고딕'
                run.font.size = Pt(11)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        # 구분선
        elif line.strip() == '---':
            doc.add_paragraph()  # 빈 줄

        # 일반 문단
        elif line.strip():
            add_paragraph(doc, line)

        i += 1

    # 문서 저장
    doc.save(output_file)
    print(f"✅ Word 문서 생성 완료: {output_file}")

if __name__ == '__main__':
    md_file = r'D:\dev1\docs\FINAL_EVALUATION_REPORT.md'
    output_file = r'D:\dev1\docs\FANS_프로젝트_평가보고서.docx'

    try:
        convert_markdown_to_word(md_file, output_file)
    except Exception as e:
        print(f"❌ 오류 발생: {e}")
        import traceback
        traceback.print_exc()
